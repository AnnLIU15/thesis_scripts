#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
周报邮件自动化工具

用法:
    python weekly_email.py receive [--date YYYYMMDD]    # 接收并下载周报
    python weekly_email.py send [--date YYYYMMDD]       # 合并并发送周报
"""

import os
import sys
import argparse
import re
import io
import glob
import shutil
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timedelta, timezone
import imaplib
import smtplib
import email as email_lib
from email.header import decode_header
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import time
import yaml
import pandas as pd
import chardet
import comtypes.client
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import qn

from PyPDF2 import PdfReader, PdfMerger, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PyPDF2.generic import AnnotationBuilder
import sxtwl
from bs4 import BeautifulSoup

import pypandoc

# 切换到脚本所在目录（确保相对路径从脚本目录计算）
script_dir = os.path.dirname(os.path.abspath(__file__))
os.chdir(script_dir)

# ============ 配置加载 ============

def load_config(config_path='config.yaml'):
    """加载并管理config.yaml配置"""
    if not os.path.exists(config_path):
        raise FileNotFoundError(f"配置文件不存在: {config_path}")

    with open(config_path, 'r', encoding='utf-8') as f:
        config = yaml.safe_load(f)

    return Config(config)


class Config:
    """配置类"""

    def __init__(self, config_dict):
        self._config = config_dict

    @property
    def accounts(self):
        return self._config.get('accounts', [])

    @property
    def recipients(self):
        return self._config.get('recipients', [])

    @property
    def cc(self):
        return self._config.get('cc', [])

    @property
    def names(self):
        return self._config.get('names', [])

    @property
    def paths(self):
        return self._config.get('paths', {})

    @property
    def download(self):
        return self._config.get('download', {})

    @property
    def pdf(self):
        return self._config.get('pdf', {})

    @property
    def threading(self):
        return self._config.get('threading', {})

    @property
    def base_path(self):
        return self.paths.get('base_path', 'E:/CourseData')

    @property
    def desktop_path(self):
        return self.paths.get('desktop_path', 'C:/Users/yourname/Desktop')

    @property
    def report_prefix(self):
        return self.paths.get('report_prefix', '周报_XXX组_')

    @property
    def attachment_keywords(self):
        return self.download.get('attachment_keywords', '周报')

    @property
    def allowed_extensions(self):
        return self.download.get('allowed_extensions', ['pdf', 'docx', 'txt', 'doc', 'md'])

    @property
    def max_workers(self):
        return self.threading.get('max_workers', 4)

    @property
    def threading_enabled(self):
        return self.threading.get('enabled', True)


# ============ 工具函数 ============

chinese_number_dict = {
    0: '零', 1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
    6: '六', 7: '七', 8: '八', 9: '九', 10: '十', 11: '十一',
    12: '十二', 13: '十三', 14: '十四', 15: '十五', 16: '十六',
    17: '十七', 18: '十八', 19: '十九', 20: '二十',
    21: '二十一', 22: '二十二', 23: '二十三', 24: '二十四',
    25: '二十五', 26: '二十六', 27: '二十七', 28: '二十八',
    29: '二十九', 30: '三十', 31: '三十一', 32: '三十二',
    33: '三十三', 34: '三十四', 35: '三十五', 36: '三十六',
    37: '三十七', 38: '三十八', 39: '三十九', 40: '四十'}

jqmc = ["冬至", "小寒", "大寒", "立春", "雨水", "惊蛰", "春分", "清明", "谷雨", "立夏",
        "小满", "芒种", "夏至", "小暑", "大暑", "立秋", "处暑", "白露", "秋分", "寒露", "霜降",
        "立冬", "小雪", "大雪"]


def chinese_number(n):
    """数字转中文"""
    return chinese_number_dict.get(n, str(n))


def get_recent_friday(date=None):
    """获取最近的星期五"""
    if date is None:
        today = datetime.now(timezone(timedelta(hours=8))).date()
    else:
        today = datetime.strptime(date, "%Y%m%d").date()

    if today.weekday() == 4:  # 星期五
        return today

    days_until_last_friday = (today.weekday() - 4) % 7
    days_until_next_friday = (4 - today.weekday() + 7) % 7

    if days_until_last_friday <= days_until_next_friday:
        return today - timedelta(days=days_until_last_friday)
    else:
        return today + timedelta(days=days_until_next_friday)


def get_date_range(friday_date):
    """获取邮件搜索日期范围"""
    three_days_ago = friday_date - timedelta(days=3)
    sun_day = friday_date + timedelta(days=2)
    six_days_ago = friday_date - timedelta(days=6)
    return {
        'start': three_days_ago.strftime('%Y%m%d'),
        'end': sun_day.strftime('%Y%m%d'),
        'friday': friday_date.strftime('%Y%m%d'),
        'six_days_ago': six_days_ago.strftime('%Y%m%d')
    }


def get_semester(date):
    """判断学期"""
    month = date.month
    if 7 <= month <= 10:
        return "秋"
    elif 1 <= month <= 3:
        return "春"
    else:
        raise ValueError(f"未知学期: {date}")


def get_jieqi(date):
    """获取节气"""
    cur_day_4_sxtwl = sxtwl.fromSolar(date.year, date.month, date.day)
    next_day_4_sxtwl = cur_day_4_sxtwl.after(1)

    if cur_day_4_sxtwl.hasJieQi():
        return f"{jqmc[cur_day_4_sxtwl.getJieQi()]}安康!"
    elif next_day_4_sxtwl.hasJieQi():
        return f"{jqmc[next_day_4_sxtwl.getJieQi()]}安康!"
    return ""


def get_time_period():
    """获取时间段（上午/中午/下午/晚上）"""
    hour = datetime.now(timezone(timedelta(hours=8))).hour
    if 0 <= hour < 12:
        return '上午'
    elif 12 <= hour < 13:
        return '中午'
    elif 13 <= hour < 18:
        return '下午'
    else:
        return '晚上'


def parse_date_folder(folder_name, prefix):
    """解析文件夹名称获取日期"""
    try:
        return datetime.strptime(folder_name, f"{prefix}%Y%m%d")
    except ValueError:
        return None


def get_week_number(prefix):
    """获取当前周数"""
    matching_folders = glob.glob(f"{prefix}*")
    if not matching_folders:
        return 1

    dates = []
    for folder in matching_folders:
        date = parse_date_folder(folder, prefix)
        if date:
            dates.append(date)

    if not dates:
        return len(matching_folders) + 1

    return len(dates)


# ============ PDF 处理 ============

def docx_to_pdf(doc_path):
    """Word转PDF（使用COM）"""
    try:
        # 初始化COM（多线程环境必须）
        import pythoncom
        pythoncom.CoInitialize()

        word = comtypes.client.CreateObject('Word.Application')
        current_directory = os.getcwd()
        doc_absolute_path = os.path.abspath(doc_path)
        pdf_path = doc_absolute_path[:-len(doc_absolute_path.split('.')[-1])] + 'pdf'

        doc = word.Documents.Open(doc_absolute_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()

        # 清理COM
        pythoncom.CoUninitialize()

        return pdf_path
    except Exception as e:
        print(f"PDF转换失败: {doc_path}, 错误: {e}")
        return None


def convert_files_parallel(file_list, max_workers=4):
    """多线程转换文件"""
    if not file_list:
        return

    results = []
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(docx_to_pdf, f): f for f in file_list}
        for future in as_completed(futures):
            file_path = futures[future]
            try:
                result = future.result()
                results.append((file_path, result))
            except Exception as e:
                print(f"转换异常: {file_path}, 错误: {e}")
                results.append((file_path, None))
    return results


def merge_pdfs_with_toc(pdf_files, output_path, add_links=True, all_names=None):
    """合并PDF并添加目录页"""
    if not pdf_files:
        print("没有PDF文件需要合并")
        return None, []

    # 获取每个PDF文件的页数和提交人
    pdf_info = []
    submitted_names = []

    for pdf_file in pdf_files:
        if not os.path.exists(pdf_file):
            continue

        name = None
        for cur_name in all_names or []:
            if cur_name in pdf_file:
                name = cur_name
                break

        if name:
            submitted_names.append(name)

        try:
            with open(pdf_file, 'rb') as f:
                reader = PdfReader(f)
                num_pages = len(reader.pages)
                pdf_info.append((pdf_file, num_pages, name))
        except Exception as e:
            print(f"读取PDF失败: {pdf_file}, 错误: {e}")

    if not pdf_info:
        print("没有有效的PDF文件")
        return None, []

    # 按页数排序
    pdf_info.sort(key=lambda x: x[1])

    # 注册中文字体
    try:
        pdfmetrics.registerFont(TTFont('YaHei', 'msyh.ttc'))
    except Exception as e:
        print(f"字体加载失败: {e}")

    # 创建目录页
    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=A4)
    can.setFont("YaHei", 18)

    can.drawString(72, 720, "周报文件顺序 | 点击文字有跳转 (每个页左上角有回跳)：")
    cur_page = 2

    for i, (pdf_file, num_pages, _) in enumerate(pdf_info, start=1):
        next_page = cur_page + num_pages
        prefix_name = f"{i}. {os.path.basename(pdf_file)}"
        postfix_name = f" (Page {cur_page}-{next_page-1})"

        text_width = pdfmetrics.stringWidth(postfix_name, "YaHei", 18)
        x_pos = A4[0] - text_width - 72

        can.drawString(72, 720 - 60 * i, prefix_name[:-4])
        can.drawString(x_pos, 720 - 60 * i, postfix_name)
        cur_page = next_page

    can.save()
    packet.seek(0)
    toc_pdf = PdfReader(packet)

    # 合并PDF
    merger = PdfMerger()
    merger.append(toc_pdf)

    for pdf_file, _, _ in pdf_info:
        try:
            merger.append(pdf_file)
        except Exception as e:
            print(f"合并PDF失败: {pdf_file}, 错误: {e}")

    merger.write(output_path)
    merger.close()
    print(f"合并PDF已保存: {output_path}")

    # 添加链接
    if add_links:
        add_navigation_links(output_path, pdf_info, toc_page_count=len(toc_pdf.pages))

    return output_path, submitted_names


def add_navigation_links(pdf_path, pdf_info, toc_page_count=1):
    """添加导航链接"""
    try:
        reader = PdfReader(open(pdf_path, 'rb'))
        writer = PdfWriter()

        for page in reader.pages:
            writer.add_page(page)

        num_of_pages = len(writer.pages)
        x1, y1, x2, y2 = writer.pages[0].mediabox
        cur_page = 1

        for i, (_, num_pages, _) in enumerate(pdf_info, start=1):
            next_page = cur_page + num_pages
            annotation = AnnotationBuilder.link(
                rect=(72, 720 - 60 * i + 20, x2 - 60, 720 - 60 * i - 10),
                target_page_index=cur_page,
            )
            writer.add_annotation(page_number=0, annotation=annotation)
            cur_page = next_page

        for idx, val in enumerate(writer.pages):
            if idx == 0:
                continue
            x1_t, y1_t, x2_t, y2_t = writer.pages[idx].mediabox
            annotation = AnnotationBuilder.free_text(
                "Back to contents",
                rect=(20, y2_t - 5, 170, y2_t - 27),
                font="Microsoft Yahei",
                bold=True,
                italic=True,
                font_size="20pt",
            )
            writer.add_annotation(page_number=idx, annotation=annotation)
            link_annotation = AnnotationBuilder.link(
                rect=(20, y2_t - 5, 170, y2_t - 27),
                target_page_index=0,
            )
            writer.add_annotation(page_number=idx, annotation=link_annotation)

        with open(pdf_path, 'wb') as link_pdf:
            writer.write(link_pdf)
    except Exception as e:
        print(f"添加导航链接失败: {e}")


def create_submission_table_html(submitted_names, all_names):
    """生成提交情况HTML表格"""
    df = pd.DataFrame()
    for name in all_names:
        df[name] = ""

    for name in submitted_names:
        df.loc["提交情况", name] = "✓"

    df = df.fillna("")
    return df.to_html(classes="table table-striped")


# ============ 邮件接收类 ============

class EmailReceiver:
    """邮件接收类"""

    def __init__(self, config):
        self.config = config
        self.all_name = config.names
        self.keywords = config.attachment_keywords
        self.allowed_extensions = config.allowed_extensions

    def run(self, date=None):
        """执行接收流程"""
        friday_date = get_recent_friday(date)
        date_range = get_date_range(friday_date)
        last_friday_date = date_range['friday']

        print(f"处理日期: {last_friday_date}")
        print(f"日期范围: {date_range['start']} - {date_range['end']}")

        # 创建输出文件夹
        output_folder = f"./{self.config.report_prefix}{last_friday_date}/"
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        # 查找或创建周文件夹
        this_week_name = self._get_or_create_week_folder(last_friday_date)
        output_folder = f"./{this_week_name}/"
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        # 并行执行：邮箱下载 + 本地MD处理
        all_doc_files = []

        with ThreadPoolExecutor(max_workers=4) as executor:
            # 任务1：下载邮件附件
            if len(self.config.accounts) > 1:
                future_download = executor.submit(
                    self._process_accounts_parallel,
                    date_range, output_folder, last_friday_date
                )
            else:
                future_download = executor.submit(
                    self._process_single_account,
                    self.config.accounts[0], date_range, output_folder, last_friday_date
                )

            # 任务2：处理本地MD文件（独立执行）
            future_md = executor.submit(
                self._process_personal_report,
                friday_date, last_friday_date, this_week_name
            )

            # 收集下载的doc文件
            try:
                doc_files = future_download.result()
                if doc_files:
                    all_doc_files.extend(doc_files)
            except Exception as e:
                print(f"邮件下载失败: {e}")

            # 等待MD处理完成
            try:
                future_md.result()
            except Exception as e:
                print(f"MD处理失败: {e}")

        # 多线程转换PDF（根据文件数量决定线程数）
        if all_doc_files:
            num_files = len(all_doc_files)
            workers = min(num_files, 8)  # 最多8个线程
            print(f"开始转换 {num_files} 个文件...")
            convert_files_parallel(all_doc_files, workers)
            print("PDF转换完成")

        print('接收完成!')

    def _get_or_create_week_folder(self, last_friday_date):
        """获取或创建周文件夹"""
        prefix = self.config.report_prefix
        matching_folders = glob.glob(f"{prefix}*")
        num_of_week = len(matching_folders)

        if num_of_week == 0:
            return f'{prefix}{last_friday_date}'

        date_list = []
        for folder in matching_folders:
            date = parse_date_folder(folder, prefix)
            if date:
                date_list.append((folder, date))

        if not date_list:
            return f'{prefix}{last_friday_date}'

        max_date = max(date_list, key=lambda x: x[1])
        if max_date[1].strftime("%Y%m%d") == last_friday_date:
            return max_date[0]
        else:
            return f'{prefix}{last_friday_date}'

    def _process_accounts_parallel(self, date_range, output_folder, last_friday_date):
        """并行处理多个账户"""
        all_doc_files = []
        with ThreadPoolExecutor(max_workers=min(self.config.max_workers, len(self.config.accounts))) as executor:
            futures = {
                executor.submit(
                    self._process_single_account,
                    account, date_range, output_folder, last_friday_date
                ): account
                for account in self.config.accounts
            }
            for future in as_completed(futures):
                account = futures[future]
                try:
                    doc_files = future.result()
                    all_doc_files.extend(doc_files)
                    print(f"账户 {account['email']} 处理完成")
                except Exception as e:
                    print(f"账户 {account['email']} 处理失败: {e}")
        return all_doc_files

    def _process_single_account(self, account, date_range, output_folder, last_friday_date):
        """处理单个账户，返回下载的doc文件列表"""
        print(f"正在处理账户: {account['email']}")
        doc_files = []

        try:
            imap_conn = imaplib.IMAP4_SSL(account['imap_server'], 993)
            imap_conn.login(account['email'], account['password'])
            print(f"登录成功: {account['email']}")

            imap_conn.select("INBOX")
            resp, mails = imap_conn.search(None, 'ALL')
            msg_list = mails[0].split()

            # 倒序处理（从最新邮件开始）
            for i in reversed(range(len(msg_list))):
                try:
                    resp, data = imap_conn.fetch(msg_list[i], '(RFC822)')
                    emailbody = data[0][1]
                    mail = email_lib.message_from_bytes(emailbody)

                    mail_date = self._parse_mail_date(mail)
                    if not mail_date:
                        continue

                    mail_date_str = mail_date.strftime("%Y%m%d")

                    if mail_date_str > date_range['end']:
                        continue
                    elif mail_date_str < date_range['start']:
                        break

                    subject = self._decode_subject(mail)
                    if self._is_valid_report(mail, subject):
                        print(f"  [{mail_date_str}] {subject}")
                        downloaded_docs = self._download_attachments_with_body(mail, subject, output_folder, last_friday_date)
                        doc_files.extend(downloaded_docs)

                except Exception as e:
                    print(f"处理邮件失败: {e}")
                    continue

            imap_conn.close()
            imap_conn.logout()

        except Exception as e:
            print(f"账户处理失败: {account['email']}, 错误: {e}")

        return doc_files

    def _parse_mail_date(self, mail):
        """解析邮件日期"""
        try:
            date_str = mail.get("Date")
            if not date_str:
                return None

            for fmt in [
                '%a, %d %b %Y %H:%M:%S',
                '%d %b %Y %H:%M:%S +0800',
                '%d %b %Y %H:%M:%S -0500',
                '%d %b %Y %H:%M:%S +0000',
            ]:
                try:
                    return datetime.strptime(date_str[:24], fmt).replace(tzinfo=timezone(timedelta(hours=8)))
                except ValueError:
                    continue

            # 尝试通用格式
            return datetime.strptime(date_str[:19], '%d %b %Y %H:%M:%S').replace(tzinfo=timezone(timedelta(hours=8)))
        except Exception:
            return None

    def _decode_subject(self, mail):
        """解码邮件主题"""
        subject = mail.get('Subject', '')
        if not subject:
            return ''

        return self._decode_header(subject)

    def _is_valid_report(self, mail, subject):
        """检查是否为有效的周报"""
        # 如果主题为空，检查发件人是否在名单中
        if not subject:
            from_addr = mail.get('From', '')
            for name in self.all_name:
                if name in from_addr:
                    return True
            return False

        name_exist = any(name in subject for name in self.all_name)
        return '周报' in subject and name_exist

    def _download_attachments(self, mail, subject, save_dir, last_friday_date):
        """下载邮件附件，返回需要转换的doc文件列表"""
        name = None
        doc_files = []

        for cur_name in self.all_name:
            if cur_name in subject:
                name = cur_name
                break

        for part in mail.walk():
            content_type = part.get_content_type()
            if content_type == 'multipart/mixed':
                continue
            content_disposition = str(part.get('Content-Disposition', ''))
            if not content_disposition:
                continue

            filename = part.get_filename()
            if not filename:
                continue

            # 确保filename是字符串
            if isinstance(filename, bytes):
                try:
                    filename = decode_header(filename)[0][0].decode('utf-8')
                except:
                    continue
            else:
                filename = str(filename)

            if self.keywords not in filename:
                continue

            extension = os.path.splitext(filename)[1].lstrip('.').lower()
            if extension not in self.allowed_extensions:
                continue

            # 确定保存名称
            save_name = name
            if save_name == 'extract':
                save_name = None
                for cur_name in self.all_name:
                    if cur_name in filename:
                        save_name = cur_name
                        break

            if not save_name:
                continue

            new_filename = f"周报_{save_name}_{last_friday_date}.{extension}"
            file_path = os.path.join(save_dir, new_filename)

            # 检查是否已有对应的PDF文件（已转换完成）
            pdf_path = file_path.rsplit('.', 1)[0] + '.pdf'
            if os.path.exists(pdf_path):
                print(f"    PDF已存在，跳过: {os.path.basename(pdf_path)}")
                continue

            # 保存文件（覆盖旧版本，保留最新）
            try:
                with open(file_path, 'wb') as f:
                    f.write(part.get_payload(decode=True))
                print(f"    下载: {new_filename}")

                # 收集需要转换的doc文件
                if extension in ['doc', 'docx']:
                    doc_files.append(file_path)
            except Exception as e:
                print(f"    保存失败: {new_filename}, 错误: {e}")

        return doc_files

    def _download_attachments_with_body(self, mail, subject, save_dir, last_friday_date):
        """下载邮件附件（含正文转换），返回需要转换的doc文件列表
        优先级：附件 > 正文
        - 有附件时，只处理附件（docx/md转pdf，pdf直接用）
        - 无附件时，才转换正文
        """
        name = None

        # 如果主题为空，尝试从发件人提取姓名
        if not subject:
            from_addr = str(mail.get('From', ''))
            for cur_name in self.all_name:
                if cur_name in from_addr:
                    name = cur_name
                    break
        else:
            for cur_name in self.all_name:
                if cur_name in subject:
                    name = cur_name
                    break

        doc_files = []
        attachment_files = []

        for part in mail.walk():
            # 跳过 multipart (容器)
            if part.get_content_maintype() == 'multipart':
                continue

            # 只处理有 Content-Disposition 的部分（附件）
            content_disposition = part.get('Content-Disposition')
            if content_disposition is None:
                continue

            filename = part.get_filename()
            if not filename:
                continue

            # 解码文件名
            try:
                decoded = decode_header(filename)
                filename = decoded[0][0].decode(decoded[0][1] or 'utf-8')
            except:
                pass

            if not filename:
                continue

            # 检查是否是周报附件
            if self.keywords not in filename:
                print(f'    附件不包含关键词，跳过: {filename}')
                continue

            extension = os.path.splitext(filename)[1].lstrip('.').lower()
            if extension not in self.allowed_extensions:
                continue

            # 确定姓名（从主题或文件名中提取）
            save_name = name
            if save_name == 'extract' or not save_name:
                for cur_name in self.all_name:
                    if cur_name in filename:
                        save_name = cur_name
                        break

            if not save_name:
                print(f"    无法识别发送人，跳过附件: {filename}")
                continue

            new_filename = f"周报_{save_name}_{last_friday_date}.{extension}"
            file_path = os.path.join(save_dir, new_filename)

            # 检查是否已有对应的PDF文件
            if extension in ['doc', 'docx', 'md']:
                pdf_path = file_path.rsplit('.', 1)[0] + '.pdf'
                if os.path.exists(pdf_path):
                    print(f"    PDF已存在，跳过: {os.path.basename(pdf_path)}")
                    continue

            # 保存文件
            if not os.path.exists(file_path):
                try:
                    with open(file_path, 'wb') as f:
                        f.write(part.get_payload(decode=True))
                    print(f"    下载: {new_filename}")
                except Exception as e:
                    print(f"    保存失败: {new_filename}, 错误: {e}")
                    continue

            # 记录需要转换的文件
            if extension in ['doc', 'docx', 'md']:
                doc_files.append(file_path)
            else:
                attachment_files.append(file_path)

        # 只有在没有附件时才转换正文
        if not attachment_files and not doc_files:
            if name:
                save_path = save_dir if save_dir else f"./{self.config.report_prefix}{last_friday_date}/"
                os.makedirs(save_path, exist_ok=True)
                body_doc_files = self._convert_email_body_to_pdf(mail, name, name, last_friday_date, save_path)
                doc_files.extend(body_doc_files)
            else:
                # 尝试从发件人获取姓名
                from_addr = str(mail.get('From', ''))
                for cur_name in self.all_name:
                    if cur_name in from_addr:
                        save_path = save_dir if save_dir else f"./{self.config.report_prefix}{last_friday_date}/"
                        os.makedirs(save_path, exist_ok=True)
                        body_doc_files = self._convert_email_body_to_pdf(mail, cur_name, cur_name, last_friday_date, save_path)
                        doc_files.extend(body_doc_files)
                        break

        return doc_files

    def _convert_email_body_to_pdf(self, mail, name, save_name, last_friday_date, save_path=None):
        """将邮件正文转换为PDF（解析HTML提取纯文本）"""
        doc_files = []
        text_content = ""
        save_path = save_path or f"./{self.config.report_prefix}{last_friday_date}/"

        for part in mail.walk():
            content_type = part.get_content_type()
            if content_type in ['text/plain', 'text/html']:
                payload = part.get_payload(decode=True)
                if payload is not None:
                    if content_type == 'text/html':
                        # 解析HTML提取纯文本
                        text_content = self._extract_text_from_html(payload)
                    else:
                        text_content = self._decode_text_payload(payload)

        if not text_content:
            return doc_files

        # 创建Word文档
        doc = Document()
        styles = doc.styles
        new_style = styles.add_style('NewStyle', WD_STYLE_TYPE.PARAGRAPH)
        new_style.font.name = 'Times New Roman'
        new_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        new_style.font.size = Pt(14)

        # 添加表格（姓名）
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = '姓名'
        self._set_cell_font(table.rows[0].cells[0], 'Times New Roman', 16, '宋体')
        table.rows[0].cells[1].text = name
        self._set_cell_font(table.rows[0].cells[1], 'Times New Roman', 16, '宋体')

        doc.add_paragraph(text_content, style='NewStyle')

        word_filename = f"周报_{save_name}_{last_friday_date}.docx"
        word_filepath = os.path.join(save_path, word_filename)
        doc.save(word_filepath)
        doc_files.append(word_filepath)

        print(f"    从正文创建: {word_filename}")

        return doc_files

    def _extract_text_from_html(self, payload):
        """从HTML中提取纯文本"""
        try:
            html_text = self._decode_text_payload(payload)
            soup = BeautifulSoup(html_text, 'html.parser')

            # 移除脚本和样式
            for tag in soup(['script', 'style', 'head', 'iframe']):
                tag.decompose()

            # 获取文本
            text = soup.get_text(separator='\n')

            # 清理空白行
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)

            return text
        except ImportError:
            # 如果没有beautifulsoup，直接解码
            return self._decode_text_payload(payload)
        except Exception as e:
            print(f"    HTML解析失败: {e}")
            return self._decode_text_payload(payload)

    def _decode_text_payload(self, payload):
        """尝试多种编码解码文本，处理邮件乱码问题"""
        if not payload:
            return ""

        # 编码检测顺序（从最可能到备用）
        encodings_to_try = []

        # 首先尝试 chardet 检测
        detected = chardet.detect(payload)
        if detected['confidence'] > 0.7 and detected['encoding']:
            encodings_to_try.append(detected['encoding'])

        # 添加常见中文编码
        encodings_to_try.extend(['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin1', 'iso-8859-1', 'cp1252'])

        # 去重
        encodings_to_try = list(dict.fromkeys(encodings_to_try))

        for encoding in encodings_to_try:
            try:
                text = payload.decode(encoding)
                # 检查是否包含无效字符
                if '\x00' in text or text.count('�') / len(text) < 0.1:
                    return text
            except (UnicodeDecodeError, LookupError):
                continue

        # 最后尝试忽略错误
        return payload.decode('utf-8', errors='ignore')

    def _decode_header(self, header_value):
        """安全解码邮件头部"""
        if not header_value:
            return ""

        try:
            decoded_parts = decode_header(header_value)
            result = ""
            for part, encoding in decoded_parts:
                if isinstance(part, bytes):
                    # 优先使用检测到的编码
                    enc = encoding or 'utf-8'
                    try:
                        result += part.decode(enc)
                    except LookupError:
                        result += part.decode('utf-8', errors='ignore')
                else:
                    result += str(part)
            return result
        except Exception:
            return str(header_value)

    def _set_cell_font(self, cell, font_name, font_size, chinese_font_name=None):
        """设置单元格字体"""
        paragraph = cell.paragraphs[0]
        run = paragraph.runs
        if len(run) == 0:
            run = paragraph.add_run()
        else:
            run = run[0]
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if chinese_font_name:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font_name)

    def _process_personal_report(self, friday_date, last_friday_date, week_folder):
        """处理个人周报（桌面MD文件）"""
        file_dir = self.config.desktop_path
        prefix = "周报"
        name = "XXXX"

        md_files = glob.glob(os.path.join(file_dir, '*.md'))
        md_files = [f for f in md_files if prefix in f and name in f]

        if not md_files:
            print("没有找到个人周报文件")
            return

        this_week = get_week_number(self.config.report_prefix)
        six_days_ago = (friday_date - timedelta(days=6)).strftime("%Y%m%d")

        for md_file_path in md_files:
            try:
                with open(md_file_path, "r", encoding="utf-8") as f:
                    content = f.read()

                # 更新周数
                week_pattern = r"\|\s+周数\s+\|\s+(.*?)\s+\|"
                matches = re.findall(week_pattern, content, re.DOTALL)
                if matches:
                    replace_str = f'第{chinese_number(this_week)}周 ({six_days_ago}~{last_friday_date})'
                    content = content.replace(matches[0], replace_str)

                # 转换中文标点
                chinese_punct = {
                    '‘': "'", '’': "'", '"': '"', '"': '"',
                    '。': '. ', '（': ' (', '）': ') ', '；': '; '
                }
                content = re.sub(r'[''""。（）；]', lambda x: chinese_punct.get(x.group(0), x.group(0)), content)

                # 更新周数中文
                week_num_pattern = r'第[一二三四五六七八九十百千万\d]+周'
                content = re.sub(week_num_pattern, f'第{chinese_number(this_week)}周', content)

                # 重命名文件
                new_md_file = f"第{chinese_number(this_week)}周周报_{name}_{last_friday_date}.md"
                new_md_path = os.path.join(file_dir, new_md_file)

                os.rename(md_file_path, new_md_path)
                with open(new_md_path, "w", encoding="utf-8") as f:
                    f.write(content)

                print(f"  个人周报: {new_md_file}")

                # 转换PDF
                pdf_file = f"{prefix}_{name}_{last_friday_date}.pdf"
                pdf_path = os.path.join(file_dir, pdf_file)

                try:
                    pypandoc.convert_file(
                        new_md_path, 'pdf',
                        outputfile=pdf_path,
                        extra_args=['--pdf-engine=xelatex', '--template=pandoc-template.latex']
                    )
                    # 移动到周报文件夹
                    dest = f"{week_folder}/{pdf_file}"
                    shutil.move(pdf_path, dest)
                    print(f"    转换PDF: {pdf_file}")
                except Exception as e:
                    print(f"    PDF转换失败: {e}")

            except Exception as e:
                print(f"处理个人周报失败: {md_file_path}, 错误: {e}")


# ============ 邮件发送类 ============

class EmailSender:
    """邮件发送类"""

    def __init__(self, config):
        self.config = config
        self.all_name = config.names
        self.recipients = config.recipients

    def run(self, date=None):
        """执行发送流程"""
        if date:
            cur_date = datetime.strptime(date, "%Y%m%d")
        else:
            # 获取最近的周报文件夹
            prefix = self.config.report_prefix
            matching_folders = glob.glob(f"{prefix}*")

            if not matching_folders:
                print("没有找到周报文件夹")
                return

            dates = []
            for folder in matching_folders:
                d = parse_date_folder(folder, prefix)
                if d:
                    dates.append((folder, d))

            if not dates:
                print("无法解析周报文件夹日期")
                return

            cur_date = max(dates, key=lambda x: x[1])[1]

        print(f"处理日期: {cur_date.strftime('%Y%m%d')}")

        # 查找周报文件夹
        folder_name = f"{self.config.report_prefix}{cur_date.strftime('%Y%m%d')}"
        if not os.path.exists(folder_name):
            print(f"文件夹不存在: {folder_name}")
            return

        # 获取PDF文件
        pdf_files = glob.glob(os.path.join(folder_name, "*.pdf"))
        if not pdf_files:
            print("没有找到PDF文件")
            return

        # 合并PDF
        output_filename = f"{folder_name}/周报_算法软件组_{cur_date.strftime('%Y%m%d')}.pdf"
        merged_pdf, submitted_names = merge_pdfs_with_toc(
            pdf_files, output_filename, add_links=True, all_names=self.all_name
        )

        if not merged_pdf:
            print("PDF合并失败")
            return

        # 生成邮件内容
        num_of_week = get_week_number(self.config.report_prefix)
        semester = get_semester(cur_date)
        seven_days_ago = cur_date - timedelta(days=6)

        time_period = get_time_period()
        jq = get_jieqi(cur_date)

        html_table = create_submission_table_html(submitted_names, self.all_name)

        message = (
            f"X老师、X老师:<br><br>{time_period}好! {jq}<br><br>"
            f"{cur_date.year}{semester}季学期第{num_of_week}周"
            f'({seven_days_ago.strftime("%Y%m%d")}-{cur_date.strftime("%Y%m%d")}) '
            f"算法软件组收集到的周报 ({len(submitted_names)}/{len(self.all_name)})，"
            f"提交情况如表所示：<br><br>{html_table}<br>"
            "具体文件可见于本邮件附件，请审阅！<br><br>XXXX"
        )

        print(f"提交情况: {len(submitted_names)}/{len(self.all_name)}")

        # 显示邮件预览
        subject = f"{self.config.report_prefix}{cur_date.strftime('%Y%m%d')}"
        print("\n" + "=" * 50)
        print("邮件预览:")
        print(f"  主题: {subject}")
        print(f"  收件人: {', '.join(self.recipients)}")
        if self.config.cc:
            print(f"  抄送: {', '.join(self.config.cc)}")
        print(f"  附件: {os.path.basename(merged_pdf)}")
        print(f"  正文: {message[:200]}..." if len(message) > 200 else f"  正文: {message}")
        print("=" * 50 + "\n")

        # 发送邮件
        self._send_email(message, merged_pdf)

        print('发送完成!')
        time.sleep(10)  # 等待10秒后再退出

    def _send_email(self, message_body, attachment_path):
        """发送邮件"""
        if not self.config.accounts:
            print("没有配置邮件账户")
            return

        account = self.config.accounts[0]

        msg = MIMEMultipart()
        msg["From"] = account['email']
        msg["To"] = ", ".join(self.recipients)
        msg["Subject"] = f"{self.config.report_prefix}{datetime.now().strftime('%Y%m%d')}"

        # 添加抄送
        if self.config.cc:
            msg["Cc"] = ", ".join(self.config.cc)

        msg.attach(MIMEText(message_body, 'html'))

        # 添加附件
        try:
            with open(attachment_path, 'rb') as f:
                attachment = MIMEApplication(f.read())
            attachment.add_header('Content-Disposition', 'attachment',
                                filename=f'{os.path.basename(attachment_path)}')
            msg.attach(attachment)
        except Exception as e:
            print(f"附件读取失败: {e}")
            return

        # 发送邮件
        try:
            server = smtplib.SMTP_SSL(account['smtp_server'], account['smtp_port'], timeout=30)
            server.login(account['email'], account['password'])
            # 收件人 + 抄送人
            all_recipients = self.recipients + (self.config.cc or [])
            server.sendmail(account['email'], all_recipients, msg.as_string())
            server.quit()
            print("邮件发送成功")
        except smtplib.SMTPException as e:
            print(f"SMTP错误: {e}")
        except Exception as e:
            print(f"邮件发送失败: {e}")


# ============ 主入口 ============

def main():
    parser = argparse.ArgumentParser(
        description='周报邮件自动化工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python weekly_email.py receive              # 接收本周周报
  python weekly_email.py receive --date 20250110  # 接收指定日期周报
  python weekly_email.py send                 # 发送周报
  python weekly_email.py send --date 20250110    # 发送指定日期周报
        """
    )

    parser.add_argument('command', choices=['receive', 'send'],
                       help='receive: 接收并下载周报; send: 合并并发送周报')
    parser.add_argument('--date', help='指定日期 (YYYYMMDD)，默认本周五/最近一周')

    args = parser.parse_args()

    # 加载配置（config.yaml 在脚本目录中）
    config_path = os.path.join(script_dir, 'config.yaml')

    try:
        config = load_config(config_path)
    except FileNotFoundError:
        print(f"错误: 配置文件不存在: {config_path}")
        sys.exit(1)
    except yaml.YAMLError as e:
        print(f"错误: 配置文件解析失败: {e}")
        sys.exit(1)

    # 执行命令
    if args.command == 'receive':
        receiver = EmailReceiver(config)
        receiver.run(date=args.date)
    else:
        sender = EmailSender(config)
        sender.run(date=args.date)


if __name__ == '__main__':
    main()
