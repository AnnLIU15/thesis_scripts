import time
import imaplib
import email
import os
from datetime import datetime, timedelta, timezone, date
from email.header import decode_header
import pypandoc

# 加密解密密码的相关代码
from cryptography.fernet import Fernet
from glob import glob
import re
import shutil
import numpy as np
import docx
import chardet  # 需要安装chardet库
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt  # 用于设置字体大小的点单位
from docx.oxml.ns import qn  # 用于处理 XML 名称空间
import comtypes.client

chinese_number_dict = {
    0: '零', 1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
    6: '六', 7: '七', 8: '八', 9: '九', 10: '十', 11: '十一',
    12: '十二', 13: '十三', 14: '十四', 15: '十五', 16: '十六',
    17: '十七', 18: '十八', 19: '十九', 20: '二十',
    21: '二十一', 22: '二十二', 23: '二十三', 24: '二十四',
    25: '二十五', 26: '二十六', 27: '二十七', 28: '二十八',
    29: '二十九', 30: '三十', 31: '三十一', 32: '三十二',
    33: '三十三', 34: '三十四', 35: '三十五', 36: '三十六',
    37: '三十七', 38: '三十八', 39: '三十九', 40: '四十'}

# 获取当前脚本文件所在目录
current_script_directory = os.path.dirname(os.path.abspath(__file__))
os.chdir(current_script_directory)


def times(x, s):
    return datetime.strptime(x, f"{s}%Y%m%d")


def convert_md_to_pdf_pypandoc(input_md_path, output_pdf_path, template_path):
    pypandoc.convert_file(
        input_md_path,
        "pdf",
        outputfile=output_pdf_path,
        extra_args=["--pdf-engine=xelatex", "--template=" + template_path],
    )


def process_msg(msg):
    print(msg)


# 定义函数：解密密码
def decrypt_password(encrypted_password):
    decrypted_password = cipher_suite.decrypt(encrypted_password)
    return decrypted_password.decode("utf-8")


class IMAP_Downemail(object):
    """
    imap邮箱下载附件(腾讯企业邮箱测试通过)
    """

    def __init__(
        self,
        account,
        pwd,
        serverurl,
        savedir,
        startdate,
        enddate,
        all_name,
        keywords,
        last_friday_date,
        exts=[".xls", ".xlsx"],
    ):
        """
        init
        :param account:   邮箱账户
        :param pwd:       密码
        :param serverurl: 接收服务器地址
        :param savedir:   文件保存路径
        :param startdate: 邮件开始日期
        :param enddate:   邮件结束日期
        :param exts:      附件拓展名
        """
        self._account = account
        self._pwd = pwd
        self._serverurl = serverurl
        self._savedir = savedir
        self._startdate = startdate
        self._enddate = enddate
        self._exts = exts
        self.all_name = all_name
        self.keywords = keywords
        self.last_friday_date = last_friday_date

    def __getAtt(self, msg, name: str = "extract"):
        attachments = []
        attachments_path = []
        has_attachment = False
        for part in msg.walk():
            if part.get_content_maintype() == "multipart":
                continue
            if part.get("Content-Disposition") is None:
                continue
            has_attachment = True
            fileName = part.get_filename()
            # 如果文件名为纯数字、字母时不需要解码，否则需要解码
            try:
                fileName = decode_header(fileName)[0][0].decode(
                    decode_header(fileName)[0][1]
                )
            except:
                pass
            if not self.keywords in fileName:
                print(f"Attachment is not contain keyword: {fileName}")
            # 只获取指定拓展名的附件
            extension = os.path.splitext(os.path.split(fileName)[1])[1][1:]
            if not extension in self._exts:
                continue
            # 如果获取到了文件，则将文件保存在指定的目录下
            if fileName:
                if name == "extract":
                    for cur_name in self.all_name:
                        if cur_name in fileName:
                            name = cur_name
                            break
                new_filename = f"周报_{name}_{self.last_friday_date}.{extension}"
                filePath = os.path.join(self._savedir, new_filename)
                if not os.path.exists(filePath):
                    fp = open(filePath, "wb")
                    fp.write(part.get_payload(decode=True))
                    fp.close()
                    if extension in ["doc", "docx"]:
                        self.doc_to_pdf(
                            filePath,
                        )
                    attachments.append(fileName)
                    attachments_path.append(filePath)
                else:
                    print(f"已存在文件: {filePath}")
            return attachments, attachments_path, has_attachment

    def __getEmailattachment(self, msg, subject):
        """
        下载邮件中的附件
        """
        name = None
        for cur_name in self.all_name:
            if cur_name in subject:
                name = cur_name
                break
        attachments, attachments_path, has_attachment = self.__getAtt(msg, name=name)

        if (not has_attachment) and len(attachments) == 0:
            for part in msg.walk():
                if part.get_content_maintype() == "text":
                    payload = part.get_payload(decode=True)
                    if payload is not None:
                        # 使用chardet检测编码
                        detected_encoding = chardet.detect(payload)["encoding"]
                        try:
                            if detected_encoding is not None:
                                # 使用检测到的编码解码
                                text = payload.decode(detected_encoding)
                            else:
                                # 作为后备，使用utf-8解码，并忽略错误
                                text = payload.decode("utf-8", errors="ignore")
                            # print(text)
                        except UnicodeDecodeError:
                            pass
                            # print("解码错误")
                    else:
                        # 处理 payload 为 None 的情况
                        pass
                    # 创建Word文档并写入正文
                    # 定义一个新的段落样式

                    doc = docx.Document()
                    styles = doc.styles
                    new_style = styles.add_style("NewStyle", WD_STYLE_TYPE.PARAGRAPH)
                    new_style.font.name = "Times New Roman"
                    new_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    new_style.font.size = Pt(14)
                    table = doc.add_table(rows=1, cols=2)
                    # 设置表格样式（可选）
                    table.style = "Table Grid"
                    table.rows[0].cells[0].text = "姓名"
                    self.set_cell_font(
                        table.rows[0].cells[0], "Times New Roman", 16, "宋体"
                    )
                    table.rows[0].cells[1].text = name
                    self.set_cell_font(
                        table.rows[0].cells[1], "Times New Roman", 16, "宋体"
                    )
                    doc.add_paragraph(text, style="NewStyle")
                    word_filename = f"周报_{name}_{self.last_friday_date}.docx"
                    word_filepath = os.path.join(self._savedir, word_filename)
                    doc.save(word_filepath)
                    self.doc_to_pdf(
                        word_filepath,
                    )
                    attachments.append(word_filename)
                    attachments_path.append(word_filepath)
                    break
        return attachments, attachments_path

    def doc_to_pdf(
        self,
        doc_path,
    ):
        word = comtypes.client.CreateObject("Word.Application")
        current_directory = os.getcwd()
        # 将相对路径转换为绝对路径
        doc_absolute_path = os.path.abspath(os.path.join(current_directory, doc_path))
        pdf_path = doc_absolute_path[: -len(doc_absolute_path.split(".")[-1])] + "pdf"
        doc = word.Documents.Open(doc_absolute_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # FileFormat=17 对应于PDF格式
        doc.Close()
        word.Quit()

    def set_cell_font(self, cell, font_name, font_size, chinese_font_name=None):
        paragraph = cell.paragraphs[0]
        run = paragraph.runs
        if len(run) == 0:
            run = paragraph.add_run()
        else:
            run = run[0]
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if chinese_font_name:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese_font_name)

    def scanDown(self, process_fun=None):
        if process_fun:
            process_fun("当前邮箱：{}".format(self._account))

        # 连接到qq企业邮箱，其他邮箱调整括号里的参数
        imap_connection = imaplib.IMAP4_SSL(self._serverurl, 993)

        if process_fun:
            process_fun("身份认证...")
        try:
            # 用户名、密码，登陆
            imap_connection.login(self._account, self._pwd)
            login_success = True
        except:
            login_success = False

        if login_success:
            if process_fun:
                process_fun("邮箱{}登录成功！".format(self._account))

            # 选定一个邮件文件夹
            # 收件箱默认名称是"INBOX"
            # 可以用conn.list()查看都有哪些文件夹
            imap_connection.select("INBOX")

            # 提取文件夹中所有邮件的编号
            resp, mails = imap_connection.search(None, "ALL")

            # 邮件编号列表
            msgList = mails[0].split()

            # 从最近的邮件开始获取
            for i in reversed(range(len(msgList))):
                try:
                    resp, data = imap_connection.fetch(msgList[i], "(RFC822)")
                    emailbody = data[0][1]
                    mail = email.message_from_bytes(emailbody)
                    # 解析邮件日期
                    try:
                        mail_date = time.strptime(
                            mail.get("Date")[0:24], "%a, %d %b %Y %H:%M:%S"
                        )  # 格式化收件时间
                    except:
                        mail_date = time.strptime(
                            mail.get("Date"), "%d %b %Y %H:%M:%S +0800"
                        )  # 格式化收件时间
                    startdate = self._startdate
                    stopdate = self._enddate
                    mail_date = time.strftime("%Y%m%d", mail_date)
                    if mail_date > stopdate:
                        continue
                    elif mail_date < startdate:
                        break
                    else:
                        subject = mail["Subject"]
                        decoded_subject = email.header.decode_header(subject)[0]
                        if isinstance(decoded_subject[0], bytes):
                            subject = decoded_subject[0].decode(
                                decoded_subject[1] or "utf-8"
                            )
                        else:
                            subject = decoded_subject[0]
                        if subject == "":
                            print(mail_date, mail["From"], subject)
                            attachments, attachments_path, _ = self.__getAtt(mail)
                            for attachment, attachment_path in zip(
                                attachments, attachments_path
                            ):
                                if process_fun:
                                    process_fun(
                                        "\t已下载文件: {} -> {}".format(
                                            attachment, attachment_path
                                        )
                                    )
                        else:
                            # 如果邮件主题使用了 MIME 编码（如 UTF-8），则解码它
                            name_exist_in_list = False

                            for cur_name in self.all_name:
                                if cur_name in subject:
                                    name_exist_in_list = True
                                    break
                            if ("周报" in subject) and name_exist_in_list:
                                print(mail_date, mail["From"], subject)
                                # 获取附件
                                attachments, attachments_path = (
                                    self.__getEmailattachment(mail, subject)
                                )
                                for attachment, attachment_path in zip(
                                    attachments, attachments_path
                                ):
                                    if process_fun:
                                        process_fun(
                                            "\t已下载文件: {} -> {}".format(
                                                attachment, attachment_path
                                            )
                                        )

                except Exception as e:
                    print(e)
                    continue

            imap_connection.close()
            imap_connection.logout()
        else:
            if process_fun:
                process_fun("邮箱{}登录失败！".format(self._account))


# 需要匹配的关键字列表
keywords = "周报"
# 发送人名字列表
all_name = ["xxxxxxx"]

allowed_extensions = ["pdf", "docx", "txt", "doc", "md"]  # 可以匹配的文件格式列表

des_path = "D:/CourseData"  # 附件保存目标文件夹 # 需要自己加密一个密码 password
with open(f"{des_path}/key", "rb") as key_file:
    key = key_file.read()

# 创建一个 Fernet 密钥对象
cipher_suite = Fernet(key)
with open(f"{des_path}/password", "rb") as f:
    encrypted_password = f.read()

# 邮箱账号列表
account_list = [
    {
        "email": "xxxxxxxx@mail2.sysu.edu.cn",  # 邮箱地址
        "password": decrypt_password(encrypted_password),  # 授权密码 # 可改明文
        "server": "imap.exmail.qq.com",  # 服务器地址
    }
]

# 计算最近的星期五
utc_plus_8 = timezone(timedelta(hours=8))
current_utc_plus_8_time = datetime.now(utc_plus_8)
today = current_utc_plus_8_time.date()
# today = date(year=2023, month=12, day=29)

# 如果今天是星期五，则当前日期就是最近的星期五
if today.weekday() == 4:  # 星期五的weekday为4
    nearest_friday = today
else:
    # 计算距离上一个星期五还有多少天
    days_until_last_friday = (today.weekday() - 4) % 7

    # 计算距离下一个星期五还有多少天
    days_until_next_friday = (4 - today.weekday() + 7) % 7

    # 判断哪个星期五离当前日期更近
    if days_until_last_friday <= days_until_next_friday:
        nearest_friday = today - timedelta(days=days_until_last_friday)
    else:
        nearest_friday = today + timedelta(days=days_until_next_friday)

# 计算六天之前的日期
six_days_ago = nearest_friday - timedelta(days=6)
print(nearest_friday)
# 将结果格式化为字符串
six_days_ago_str = six_days_ago.strftime("%Y%m%d")

# 按照 yyyymmdd 格式输出日期
last_friday_date = nearest_friday.strftime("%Y%m%d")

three_days_ago_utc_plus_8 = nearest_friday - timedelta(days=3)
sun_day = nearest_friday + timedelta(days=2)
# 使用glob模块获取匹配的文件夹
prefix_folder = "周报_xxxxx组_"
matching_folders = glob(f"{prefix_folder}*")
num_of_week = len(matching_folders)
if num_of_week == 0:
    this_week = 1
    this_week_name = f"{prefix_folder}{last_friday_date}"
else:
    date_list = [times(folder, prefix_folder) for folder in matching_folders]
    last_folder_date = max(date_list)
    last_folder_date = last_folder_date.strftime("%Y%m%d")
    if last_folder_date == last_friday_date:
        this_week = num_of_week
        this_week_name = matching_folders[np.argmax(date_list)]
    else:
        this_week = num_of_week + 1
        this_week_name = f"{prefix_folder}{last_friday_date}"
print(this_week_name)
# if not this_week == num_of_week:
#
# 文件保存目录
output_folder = f"./{prefix_folder}{last_friday_date}/"  # 附件保存目标文件夹
if not os.path.exists(output_folder):
    os.makedirs(output_folder)
startdate = three_days_ago_utc_plus_8.strftime("%Y%m%d")
enddate = sun_day.strftime("%Y%m%d")
print(
    startdate,
    enddate,
    nearest_friday.strftime("%Y%m%d"),
)
# 下载
for account in account_list:
    _email = account["email"]
    _password = account["password"]
    _server = account["server"]
    etool = IMAP_Downemail(
        _email,
        _password,
        _server,
        output_folder,
        startdate,
        enddate,
        all_name,
        keywords,
        nearest_friday.strftime("%Y%m%d"),
        exts=allowed_extensions,
    )
    etool.scanDown(process_msg)
print("Done.")


# 定义前缀变量
prefix = "周报"
name = "xxxxxxx"
file_dir = "~/Desktop"
# 使用glob模块获取匹配的文件夹

md_files = glob(file_dir + "/*.md")
md_files = [val for val in md_files if ((prefix in val) and (name in val))]
for md_file_path in md_files:
    with open(md_file_path, "r", encoding="utf-8") as file:
        markdown_content = file.read()
    pattern = r"\|\s+周数\s+\|\s+(.*?)\s+\|"
    matches = re.findall(pattern, markdown_content, re.DOTALL)
    replace_str = (
        f"第{chinese_number_dict[this_week]}周 ({six_days_ago_str}~{last_friday_date})"
    )
    print(matches[0])
    if matches:
        markdown_content = markdown_content.replace(matches[0], replace_str)
        print("Revision:", matches[0], "->", replace_str)
        chinese_to_english_mapping = {
            "‘": "'",  # 中文单引号到英文单引号
            "’": "'",  # 中文单引号到英文单引号
            "“": '"',  # 中文双引号到英文双引号
            "”": '"',  # 中文双引号到英文双引号
            "。": ". ",  # 中文句号到英文句号
            "（": " (",  # 中文左括号到英文左括号
            "）": ") ",  # 中文右括号到英文右括号
            "；": "; ",  # 中文分号到英文分号
        }
        pattern = r"第[一二三四五六七八九十百千万\d]+周"
        replace_str_week = f"第{chinese_number_dict[this_week]}周"
        matches = re.findall(pattern, markdown_content, re.DOTALL)
        for match in matches:
            markdown_content = markdown_content.replace(match, replace_str_week)
            print("Revision:", match, "->", replace_str_week)
        # 使用re.sub函数将中文符号替换为英文符号
        markdown_content = re.sub(
            r"[‘’“”（）；。]",
            lambda x: chinese_to_english_mapping[x.group(0)],
            markdown_content,
        )
        # 如果需要，你可以将修改后的Markdown内容保存回文件中
        with open(md_file_path, "w", encoding="utf-8") as file:
            file.write(markdown_content)
        # 新的MD文件路径
        new_md_file = (
            f"第{chinese_number_dict[this_week]}周周报_" f"{name}_{last_friday_date}.md"
        )
        new_md_file_path = f"{file_dir}/" f"{new_md_file}"
        os.rename(md_file_path, new_md_file_path)
        print("Name:", md_file_path, "->", new_md_file_path)
        # 构建PDF文件名，包括前缀和星期五的日期
        pdf_file_name = f"{prefix}_{name}_{last_friday_date}.pdf"
        pdf_file_path = os.path.join(file_dir, pdf_file_name)
        os.chdir(file_dir)
        pandoc_command = (
            f"pandoc {new_md_file_path} "
            f"--pdf-engine=xelatex -o {pdf_file_path} "
            "--template=pandoc-template.latex"
        )
        convert_md_to_pdf_pypandoc(
            new_md_file_path, pdf_file_path, "pandoc-template.latex"
        )

        destination_file = new_md_file[:-2] + "pdf"
        current_script_directory = os.path.dirname(os.path.abspath(__file__))
        os.chdir(current_script_directory)
        print("Export:", pandoc_command)
        shutil.copy(pdf_file_path, destination_file)
        print("Copy:", pdf_file_path, "->", destination_file)
        shutil.copy(new_md_file_path, new_md_file)
        print("Copy:", new_md_file_path, "->", new_md_file)
        destination_file = f"{this_week_name}/{pdf_file_name}"
        shutil.move(pdf_file_path, destination_file)
        print("Move:", pdf_file_path, "->", destination_file)
        ## copy assets
        source_folder = os.path.join(file_dir, "assets")
        destination_folder = "./assets"
        try:
            # 尝试将源文件夹内容复制到目标文件夹中
            shutil.copytree(source_folder, destination_folder)
            print(f"Merge: {source_folder} -> {destination_folder}")
        except FileExistsError:
            # 如果目标文件夹已存在，则将源文件夹内容添加到目标文件夹中
            for item in os.listdir(source_folder):
                source_item = os.path.join(source_folder, item)
                destination_item = os.path.join(destination_folder, item)
                if os.path.isdir(source_item):
                    shutil.copytree(source_item, destination_item)
                else:
                    shutil.copy2(source_item, destination_item)
            print(f"Merge: {source_folder} -> {destination_folder}")
        except Exception as e:
            print(f"An error occurred: {e}")
